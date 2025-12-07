"""
Document parsing service.
Converts Word documents (.docx) into structured content blocks.
"""

import io
from typing import List
from docx import Document

from app.models import ContentBlock


def parse_document(file_bytes: bytes) -> List[ContentBlock]:
    """
    Parse a Word document into a list of content blocks.
    
    Args:
        file_bytes: Raw bytes of the .docx file
        
    Returns:
        List of ContentBlock objects representing document structure
    """
    doc = Document(io.BytesIO(file_bytes))
    blocks = []
    index = 0
    
    # Iterate through document body elements in order
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]  # Get tag name without namespace
        
        if tag == 'p':  # Paragraph
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if text:  # Skip empty paragraphs
                        blocks.append(ContentBlock(
                            index=index,
                            block_type="paragraph",
                            content=text
                        ))
                        index += 1
                    break
                    
        elif tag == 'tbl':  # Table
            for table in doc.tables:
                if table._element is element:
                    table_text = _table_to_text(table)
                    blocks.append(ContentBlock(
                        index=index,
                        block_type="table",
                        content=table_text
                    ))
                    index += 1
                    break
    
    return blocks


def _table_to_text(table) -> str:
    """
    Convert a Word table to comparable text format.
    
    Format:
        [TABLE]
        Header1 | Header2 | Header3
        ---
        Val1 | Val2 | Val3
        [/TABLE]
    """
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
        if i == 0:  # Add separator after header
            rows.append("---")
    
    return "[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]"
