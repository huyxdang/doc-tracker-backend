"""
Annotated document generation service.
Creates Word documents with highlighted changes and comments.
"""

import io
from typing import List

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models import ClassifiedChange, ImpactLevel


# Impact level to highlight color mapping
IMPACT_COLORS = {
    ImpactLevel.CRITICAL: RGBColor(255, 102, 102),   # Red
    ImpactLevel.MEDIUM: RGBColor(255, 255, 102),     # Yellow
    ImpactLevel.LOW: RGBColor(200, 200, 200),        # Gray
}


def create_annotated_document(
    modified_bytes: bytes,
    classified_changes: List[ClassifiedChange]
) -> bytes:
    """
    Create a copy of the modified document with:
    1. Highlighted changes (color-coded by impact)
    2. Comments explaining each change
    
    Args:
        modified_bytes: Raw bytes of the modified .docx file
        classified_changes: List of classified changes to annotate
        
    Returns:
        Bytes of the annotated .docx file
    """
    doc = Document(io.BytesIO(modified_bytes))
    block_map = _build_block_map(doc)
    
    for change in classified_changes:
        block_idx = _extract_block_index(change.location)
        
        if block_idx is None or block_idx not in block_map:
            continue
            
        element_info = block_map[block_idx]
        
        if element_info['type'] == 'paragraph':
            _annotate_paragraph(element_info['element'], change, doc)
        elif element_info['type'] == 'table':
            _annotate_table(element_info['element'], change, doc)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()


def _build_block_map(doc: Document) -> dict:
    """Build a mapping of block index -> document element."""
    block_map = {}
    index = 0
    
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        
        if tag == 'p':
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if text:
                        block_map[index] = {
                            'type': 'paragraph',
                            'element': para
                        }
                        index += 1
                    break
                    
        elif tag == 'tbl':
            for table in doc.tables:
                if table._element is element:
                    block_map[index] = {
                        'type': 'table',
                        'element': table
                    }
                    index += 1
                    break
    
    return block_map


def _extract_block_index(location: str) -> int | None:
    """Extract block index from location string like 'Block 16'."""
    import re
    match = re.search(r'Block (\d+)', location)
    if match:
        return int(match.group(1)) - 1
    return None


def _annotate_paragraph(para, change: ClassifiedChange, doc: Document):
    """Highlight a paragraph and add a comment."""
    color = IMPACT_COLORS.get(change.impact, RGBColor(255, 255, 0))
    
    for run in para.runs:
        run.font.highlight_color = None
        _set_run_highlight(run, color)
    
    comment_text = _format_comment(change)
    _add_comment_to_paragraph(para, comment_text, doc)


def _annotate_table(table, change: ClassifiedChange, doc: Document):
    """Highlight changed cells in a table and add a comment."""
    color = IMPACT_COLORS.get(change.impact, RGBColor(255, 255, 0))
    
    changed_cells = _find_changed_cells(table, change)
    
    if changed_cells:
        for row_idx, col_idx in changed_cells:
            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                _highlight_cell(cell, color)
    else:
        for row in table.rows:
            for cell in row.cells:
                _highlight_cell(cell, color)
    
    if table.rows and table.rows[0].cells:
        first_cell = table.rows[0].cells[0]
        if first_cell.paragraphs:
            comment_text = _format_comment(change)
            _add_comment_to_paragraph(first_cell.paragraphs[0], comment_text, doc)


def _find_changed_cells(table, change: ClassifiedChange) -> List[tuple]:
    """Find which cells changed based on word_changes."""
    changed_cells = []
    
    if not change.word_changes:
        return changed_cells
    
    changed_values = set()
    for wc in change.word_changes:
        if wc.old_text:
            changed_values.add(wc.old_text.strip())
        if wc.new_text:
            changed_values.add(wc.new_text.strip())
    
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            for val in changed_values:
                if val in cell_text or cell_text in val:
                    changed_cells.append((row_idx, col_idx))
                    break
    
    return changed_cells


def _highlight_cell(cell, color: RGBColor):
    """Apply background shading to a table cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            _set_run_highlight(run, color)
    _set_cell_shading(cell, color)


def _set_cell_shading(cell, color: RGBColor):
    """Set background color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), f'{color[0]:02x}{color[1]:02x}{color[2]:02x}')
    tcPr.append(shd)


def _set_run_highlight(run, color: RGBColor):
    """Set background highlight for a text run using shading."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{color[0]:02x}{color[1]:02x}{color[2]:02x}')
    rPr.append(shd)


def _format_comment(change: ClassifiedChange) -> str:
    """Format the comment text for a change (simple format)."""
    impact_vn = {
        "critical": "NGHIÊM TRỌNG",
        "medium": "TRUNG BÌNH", 
        "low": "THẤP"
    }.get(change.impact.value, change.impact.value.upper())
    
    original = change.original[:100] if change.original else '(không có)'
    if change.original and len(change.original) > 100:
        original += '...'
        
    modified = change.modified[:100] if change.modified else '(không có)'
    if change.modified and len(change.modified) > 100:
        modified += '...'
    
    return f"""[{impact_vn}]
Gốc: {original}
Mới: {modified}"""


def _add_comment_to_paragraph(para, comment_text: str, doc: Document):
    """Add a visual comment marker to a paragraph."""
    if para.runs:
        marker_run = para.add_run(" [📝 ")
        marker_run.font.size = Pt(8)
        marker_run.font.color.rgb = RGBColor(128, 128, 128)
        
        short_comment = comment_text.split('\n')[0][:50]
        comment_run = para.add_run(short_comment)
        comment_run.font.size = Pt(8)
        comment_run.font.color.rgb = RGBColor(128, 128, 128)
        comment_run.font.italic = True
        
        close_run = para.add_run("]")
        close_run.font.size = Pt(8)
        close_run.font.color.rgb = RGBColor(128, 128, 128)
